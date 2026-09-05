"""
Video-to-Notes AI Platform — STEP 2: Export
Turns the structured Markdown notes into downloadable files: .md / .docx / .pdf.

The notes always follow a predictable shape (see the prompt in Section 4 of the plan):
  # / ## headings, "- " bullet points, **bold** key terms, plain TL;DR / question paragraphs.
This module parses that shape directly rather than pulling in a full Markdown-to-anything
rendering engine, which keeps Step 2's dependency list small.
"""

import io
import re
from dataclasses import dataclass
from timeutils import parse_timestamp


# ---------------------------------------------------------------------------
# Shared line-level markdown parsing
# ---------------------------------------------------------------------------

@dataclass
class Line:
    kind: str   # 'h1' | 'h2' | 'h3' | 'bullet' | 'numbered' | 'paragraph' | 'blank'
    text: str   # content with the markdown marker stripped off


def parse_markdown(markdown: str) -> list[Line]:
    lines: list[Line] = []
    for raw in markdown.splitlines():
        stripped = raw.strip()
        if not stripped:
            lines.append(Line("blank", ""))
        elif stripped.startswith("### "):
            lines.append(Line("h3", stripped[4:]))
        elif stripped.startswith("## "):
            lines.append(Line("h2", stripped[3:]))
        elif stripped.startswith("# "):
            lines.append(Line("h1", stripped[2:]))
        elif re.match(r"^[-*]\s+", stripped):
            lines.append(Line("bullet", re.sub(r"^[-*]\s+", "", stripped)))
        elif re.match(r"^\d+[.)]\s+", stripped):
            lines.append(Line("numbered", re.sub(r"^\d+[.)]\s+", "", stripped)))
        else:
            lines.append(Line("paragraph", stripped))
    return lines


def _split_bold(text: str) -> list[tuple[str, bool]]:
    """Split 'plain **bold** plain' into [(text, is_bold), ...] segments."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    segments = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            segments.append((part[2:-2], True))
        else:
            segments.append((part, False))
    return segments


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def markdown_to_docx(markdown: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    for line in parse_markdown(markdown):
        if line.kind == "blank":
            continue
        elif line.kind == "h1":
            doc.add_heading(line.text, level=1)
        elif line.kind == "h2":
            doc.add_heading(line.text, level=2)
        elif line.kind == "h3":
            doc.add_heading(line.text, level=3)
        elif line.kind in ("bullet", "numbered"):
            style = "List Bullet" if line.kind == "bullet" else "List Number"
            p = doc.add_paragraph(style=style)
            for text, bold in _split_bold(line.text):
                run = p.add_run(text)
                run.bold = bold
        else:  # paragraph
            p = doc.add_paragraph()
            for text, bold in _split_bold(line.text):
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def markdown_to_pdf(markdown: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="H2Custom", parent=styles["Heading2"], spaceBefore=14))
    styles.add(ParagraphStyle(name="H3Custom", parent=styles["Heading3"], spaceBefore=10))
    styles.add(ParagraphStyle(name="BodyCustom", parent=styles["BodyText"], spaceAfter=8, leading=15))

    def inline(text: str) -> str:
        # reportlab Paragraph understands basic <b> tags, so convert **bold** -> <b>bold</b>
        return re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)

    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * inch)]
    bullet_buffer: list[str] = []

    def flush_bullets():
        if bullet_buffer:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(inline(t), styles["BodyCustom"])) for t in bullet_buffer],
                    bulletType="bullet",
                    leftIndent=18,
                )
            )
            bullet_buffer.clear()

    for line in parse_markdown(markdown):
        if line.kind in ("bullet", "numbered"):
            bullet_buffer.append(line.text)
            continue
        flush_bullets()
        if line.kind == "blank":
            continue
        elif line.kind == "h1":
            story.append(Paragraph(inline(line.text), styles["Heading1"]))
        elif line.kind == "h2":
            story.append(Paragraph(inline(line.text), styles["H2Custom"]))
        elif line.kind == "h3":
            story.append(Paragraph(inline(line.text), styles["H3Custom"]))
        else:
            story.append(Paragraph(inline(line.text), styles["BodyCustom"]))
    flush_bullets()

    buffer = io.BytesIO()
    SimpleDocTemplate(
        buffer, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
    ).build(story)
    return buffer.getvalue()

_TIMED_HEADING_RE = re.compile(r"^\[(?P<start>[\d:]+)\u2013(?P<end>[\d:]+)\]\s*(?P<title>.+)$")


def extract_timed_sections(markdown: str) -> list[dict]:
    """
    Scans H1/H2 headings for the `[MM:SS\u2013MM:SS] Title` pattern the mode
    skeletons instruct the model to produce (learning_modes.py). Returns one
    dict per match: {"start_seconds", "end_seconds", "heading"}.

    Headings that don't match (e.g. a note generated from pasted text, which
    has no timing at all) are skipped, not errored on -- this is a no-op for
    untimed sources, exactly like chunk_by_time() vs chunk_text() in main.py.
    """
    sections = []
    for line in parse_markdown(markdown):
        if line.kind not in ("h1", "h2"):
            continue
        match = _TIMED_HEADING_RE.match(line.text.strip())
        if not match:
            continue
        try:
            start = parse_timestamp(match.group("start"))
            end = parse_timestamp(match.group("end"))
        except ValueError:
            continue
        sections.append({
            "start_seconds": start,
            "end_seconds": end,
            "heading": match.group("title").strip(),
        })
    return sections

def split_into_sections(markdown: str) -> list[dict]:
    """
    Splits raw markdown into sections at each H1/H2 boundary. Returns
    [{"index", "heading", "markdown_text"}, ...] in order.

    IMPORTANT: builds each section's text from the ORIGINAL raw lines, not
    from parse_markdown()'s stripped-and-reconstructed Line objects.
    parse_markdown() strips markdown markers (e.g. numbered-list "1." is
    thrown away entirely, bullets normalize "*" and "-" to one shape) --
    reconstructing text from that would subtly reformat every section on
    every save, even ones nobody touched. Using parse_markdown() only to
    IDENTIFY where headings are, and slicing the original text at those
    points, keeps untouched sections byte-for-byte identical when spliced
    back in Phase 2.2's regenerate/edit endpoints.
    """
    raw_lines = markdown.splitlines()
    parsed = parse_markdown(markdown)   # same line-for-line order as raw_lines

    sections: list[dict] = []
    current_heading = None
    current_raw: list[str] = []

    def _flush():
        if current_heading is None and not current_raw:
            return
        sections.append({
            "index": len(sections),
            "heading": current_heading or "",
            "markdown_text": "\n".join(current_raw).strip("\n"),
        })

    for raw_line, line in zip(raw_lines, parsed):
        if line.kind in ("h1", "h2"):
            _flush()
            current_heading = line.text
            current_raw = [raw_line]
        else:
            current_raw.append(raw_line)

    _flush()
    return sections